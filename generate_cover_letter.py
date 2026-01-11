#!/usr/bin/env python3
"""
Generate a professional cover letter for 2degrees Senior .NET Developer position
Generic enough to adapt for other positions
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import os
from datetime import datetime

# Personal Information
PERSONAL_INFO = {
    'name': 'Waqas Ahmad',
    'title': 'Senior Software Engineer & Technical Lead',
    'location': 'Selangor, Malaysia',
    'phone': '+60146806067',
    'email': 'devwithwaqas@gmail.com',
    'linkedin': 'https://www.linkedin.com/in/waqas1430/',
    'portfolio': 'https://www.waqasahmad.com'
}

def create_cover_letter():
    """Create a professional cover letter"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Date (current date)
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
    date_run.font.size = Pt(11)
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()
    
    # Hiring Manager (generic)
    doc.add_paragraph("Dear Hiring Manager,")
    doc.add_paragraph()
    
    # Opening paragraph - Engaging and specific
    opening = doc.add_paragraph()
    opening.add_run(
        "I am writing to express my strong interest in the Senior .NET Developer position at 2degrees. "
        "With over 17 years of experience architecting and delivering enterprise-scale .NET solutions for "
        "Fortune 500 companies, I am excited about the opportunity to bring my expertise in scalable microservices, "
        "cloud-native architectures, and modern software engineering practices to your team."
    )
    opening_format = opening.paragraph_format
    opening_format.space_after = Pt(12)
    
    # Second paragraph - Technical expertise and achievements
    technical = doc.add_paragraph()
    technical.add_run(
        "Throughout my career, I have consistently delivered high-impact results: architected Azure cloud solutions "
        "that reduced infrastructure costs by 40% while achieving 65% performance improvements, led the development "
        "of mission-critical systems processing 2.5M+ data points daily, and delivered projects 25% faster than "
        "industry standards with 99.9% uptime. My expertise spans the full .NET ecosystem—from ASP.NET Core and "
        "Entity Framework to microservices architecture and RESTful API design. I have extensive experience with "
        "Azure services (App Services, Functions, Service Fabric, Cosmos DB), modern frontend frameworks (Angular, "
        "Vue.js), and DevOps practices including CI/CD automation, containerization, and infrastructure as code."
    )
    technical_format = technical.paragraph_format
    technical_format.space_after = Pt(12)
    
    # Third paragraph - Leadership and team impact
    leadership = doc.add_paragraph()
    leadership.add_run(
        "Beyond technical execution, I bring strong leadership capabilities honed through mentoring 50+ developers "
        "across global teams and leading cross-functional teams of 8+ engineers. I have a proven track record of "
        "establishing Agile/Scrum practices, standardizing development processes, and fostering collaborative "
        "environments that drive both technical excellence and business outcomes. My experience includes leading "
        "the revamp of AirAsia Super App serving millions of users, architecting smart city solutions that achieved "
        "50-60% revenue growth, and building microservices architectures integrating 7+ enterprise systems under "
        "unified authentication."
    )
    leadership_format = leadership.paragraph_format
    leadership_format.space_after = Pt(12)
    
    # Fourth paragraph - Innovation and modern practices
    innovation = doc.add_paragraph()
    innovation.add_run(
        "I am passionate about leveraging cutting-edge tools and methodologies to enhance productivity and code quality. "
        "I actively utilize AI-assisted development tools (Cursor AI, Claude, ChatGPT) to accelerate development, "
        "improve architectural decision-making, and maintain high code quality standards. This forward-thinking approach, "
        "combined with my deep understanding of SOLID principles, design patterns, and clean architecture, enables me "
        "to build maintainable, scalable, and robust systems. I am committed to continuous learning and staying current "
        "with industry best practices, which I believe is essential in a dynamic environment like 2degrees."
    )
    innovation_format = innovation.paragraph_format
    innovation_format.space_after = Pt(12)
    
    # Fifth paragraph - Why 2degrees / Why this role
    why_company = doc.add_paragraph()
    why_company.add_run(
        "I am particularly drawn to this opportunity because of 2degrees' commitment to innovation, customer-centricity, "
        "and delivering reliable, modern digital experiences. The prospect of contributing to a company that values "
        "simplifying technology while maintaining high standards of quality and performance aligns perfectly with my "
        "professional values. I am excited about the opportunity to work with a collaborative, forward-thinking team "
        "where I can both contribute my expertise and continue to grow as a technical leader. I am particularly "
        "interested in how I can help enhance your .NET infrastructure, improve system scalability, and mentor team "
        "members to raise the overall technical bar."
    )
    why_company_format = why_company.paragraph_format
    why_company_format.space_after = Pt(12)
    
    # Closing paragraph
    closing = doc.add_paragraph()
    closing.add_run(
        "I am confident that my combination of deep technical expertise, proven leadership experience, and passion for "
        "building exceptional software solutions would make me a valuable addition to your team. I would welcome the "
        "opportunity to discuss how my skills and experience can contribute to 2degrees' continued success. Please find "
        "my portfolio at "
    )
    portfolio_run = closing.add_run(PERSONAL_INFO['portfolio'])
    portfolio_run.font.color.rgb = RGBColor(5, 99, 187)
    closing.add_run(" for a detailed view of my projects and technical capabilities.")
    closing_format = closing.paragraph_format
    closing_format.space_after = Pt(12)
    
    # Final closing
    doc.add_paragraph()
    doc.add_paragraph("Thank you for considering my application. I look forward to the opportunity to discuss this role further.")
    doc.add_paragraph()
    doc.add_paragraph("Kind regards,")
    doc.add_paragraph()
    
    # Signature
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(PERSONAL_INFO['name'])
    name_run.bold = True
    name_run.font.size = Pt(12)
    
    # Contact info below name
    contact_para = doc.add_paragraph()
    contact_text = f"{PERSONAL_INFO['phone']} | {PERSONAL_INFO['email']} | {PERSONAL_INFO['linkedin']}"
    contact_run = contact_para.add_run(contact_text)
    contact_run.font.size = Pt(10)
    contact_run.italic = True
    
    # Save document
    filename = 'resumes/Waqas_Ahmad_Cover_Letter.docx'
    doc.save(filename)
    print(f"✓ Created: {filename}")
    print(f"✓ Full path: {os.path.abspath(filename)}")

if __name__ == '__main__':
    print("Generating professional cover letter...")
    print("=" * 60)
    create_cover_letter()
    print("=" * 60)
    print("Cover letter generated successfully!")


