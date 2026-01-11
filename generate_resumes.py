#!/usr/bin/env python3
"""
Generate 3 ATS-compatible Word resume documents for New Zealand job market:
1. Senior Software Engineer
2. Software Architect
3. Technical Manager/Lead
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

# Personal Information
PERSONAL_INFO = {
    'name': 'Waqas Ahmad',
    'title': 'Senior Software Engineer & Technical Lead',
    'location': 'Selangor, Malaysia (Open to Relocate to New Zealand)',
    'phone': '+60146806067',
    'email': 'devwithwaqas@gmail.com',
    'linkedin': 'https://www.linkedin.com/in/waqas1430/',
    'github': 'https://github.com/devwithwaqas',
    'portfolio': 'https://www.waqasahmad.com',
    'website': 'https://www.waqasahmad.com'
}

# Skills - Comprehensive list
SKILLS = {
    'languages': ['C#', '.NET Core', 'ASP.NET Core', 'JavaScript', 'TypeScript', 'HTML5', 'CSS3', 'SQL', 'Python', 'Java'],
    'frameworks': ['Angular', 'Vue.js', 'React', 'Bootstrap', 'jQuery', 'Entity Framework Core', 'ASP.NET MVC', 'SignalR'],
    'databases': ['SQL Server', 'PostgreSQL', 'PostGIS', 'Azure Cosmos DB', 'Redis', 'Azure Data Lake'],
    'cloud': ['Microsoft Azure', 'Azure App Services', 'Azure Functions', 'Azure Service Fabric', 'Azure Key Vault', 
              'Azure Application Insights', 'Azure Storage', 'Azure DevOps', 'Azure API Management', 'Azure Event Grid',
              'Azure Service Bus', 'Azure Stream Analytics', 'Azure Active Directory', 'OAuth 2.0'],
    'devops': ['CI/CD', 'Azure DevOps', 'Git', 'GitHub', 'Docker', 'Kubernetes', 'ARM Templates', 'Infrastructure as Code',
               'Automated Testing', 'Unit Testing', 'Integration Testing'],
    'tools': ['Visual Studio', 'VS Code', 'Swagger', 'Postman', 'JIRA', 'Confluence', 'Power Apps', 'SharePoint'],
    'methodologies': ['Agile', 'Scrum', 'Microservices Architecture', 'RESTful APIs', 'GraphQL', 'SOLID Principles', 
                      'Design Patterns', 'MVVM', 'MVC', 'Clean Architecture', 'Domain-Driven Design'],
    'ai_tools': ['Cursor AI', 'Claude AI', 'ChatGPT', 'GitHub Copilot', 'AI-Assisted Development', 'AI Code Generation'],
    'gis': ['SuperMap GIS', 'Leaflet', 'Google Maps API', 'GeoJSON', 'PostGIS', 'Spatial Data Processing'],
    'monitoring': ['Grafana', 'Prometheus', 'Application Insights', 'Performance Monitoring', 'Log Analytics'],
    'other': ['IoT Integration', 'OCR Systems', 'Real-time Analytics', 'Data Analytics', 'LINQ Optimization', 
              'Database Indexing', 'Performance Tuning', 'Code Review', 'Technical Mentoring', 'Team Leadership']
}

# Work Experience
EXPERIENCE = [
    {
        'title': 'Lead Software Engineer & Technical Consultant',
        'company': 'Independent Consultant',
        'location': 'Remote/Global',
        'period': 'January 2022 - Present',
        'duration': '2+ years',
        'achievements': [
            'Delivered enterprise solutions to Fortune 500 clients across 5+ countries, specializing in Azure cloud architecture and .NET Core microservices',
            'Architected Azure cloud solutions reducing infrastructure costs by 40% while achieving 65% performance improvements for mission-critical systems',
            'Built complete applications from database design to user interface, ensuring seamless integration and optimal user experience',
            'Delivered projects 25% faster than industry standards with 99.9% uptime for mission-critical petroleum operations managing billions in maintenance costs',
            'Pioneered remote-first development practices across 5+ time zones, managing distributed teams for Fortune 500 clients with 99.9% delivery success rate',
            'Leveraged AI-powered development tools (Cursor AI, Claude, ChatGPT) to accelerate code generation, improve code quality, and enhance architectural decision-making',
            'Implemented AI-assisted development workflows, reducing development time by 30% while maintaining high code quality standards',
            'Utilized AI tools for automated code review, documentation generation, and technical design pattern recommendations'
        ]
    },
    {
        'title': 'Technical Lead',
        'company': 'AirAsia Tech',
        'location': 'Kuala Lumpur, Malaysia',
        'period': 'June 2021 – August 2022',
        'duration': '1 year 3 months',
        'achievements': [
            'Led revamp of AirAsia Super App and airline domain applications using .NET Core + Angular, serving millions of users across Southeast Asia',
            'Standardized CI/CD pipelines, branching strategies, and deployment approvals, reducing delivery time by 40-50% while ensuring code quality',
            'Optimized Entity Framework + LINQ queries, improving database performance by 35% and reducing query execution time for real-time flight operations',
            'Led cross-functional team of 8+ engineers developing multi-tier .NET Core + Angular applications, conducting code reviews and knowledge transfer sessions',
            'Designed scalable Azure Cloud Services architecture supporting real-time flight operations'
        ]
    },
    {
        'title': 'Technical Lead',
        'company': 'Datacom Systems Sdn Bhd',
        'location': 'Kuala Lumpur, Malaysia',
        'period': 'February 2020 – January 2021',
        'duration': '1 year',
        'achievements': [
            'Developed microservices architecture with API Gateway integrating 7+ third-party systems (SharePoint, SAP, Cherwell, Success Factor, Datalake) under unified Azure AD authentication',
            'Designed Azure Service Fabric cloud infrastructure with SQL Server, implemented CI/CD pipelines with ARM templates for multi-tenant deployments across DEV/QA/UAT/PROD environments',
            'Optimized n-tier application architecture with Entity Framework best practices, improving query performance and system responsiveness by 30%',
            'Integrated enterprise systems (SAP, Cherwell HR, Application Insights) with custom APIs, enabling seamless data flow across business operations',
            'Implemented Agile/Scrum practices with sprint planning, backlog management, and daily standups, ensuring on-time delivery of all project milestones'
        ]
    },
    {
        'title': 'Technical Lead',
        'company': 'Syntronic Malaysia Sdn Bhd',
        'location': 'Kuala Lumpur, Malaysia',
        'period': 'July 2019 – January 2020',
        'duration': '7 months',
        'achievements': [
            'Developed ASP.NET Core REST APIs with Swagger documentation and SignalR for real-time bidirectional communication in IoT monitoring systems',
            'Successfully revamped major module (backend, frontend, database) in under 2 months, delivering Angular SPA with responsive design and enhanced UX',
            'Implemented CI/CD automation with Azure DevOps, Docker containerization, and automated unit testing, reducing manual testing time by 60%',
            'Utilized Cosmos DB, Azure Functions, and Azure Storage for scalable solutions',
            'Active participation in Agile ceremonies and comprehensive code reviews'
        ]
    },
    {
        'title': 'Technical Lead',
        'company': 'Iconic Control Sdn Bhd',
        'location': 'Petaling Jaya, Malaysia',
        'period': 'May 2018 – July 2019',
        'duration': '1 year 3 months',
        'achievements': [
            'Architected 7+ smart city solutions from scratch (Transportation, Law Enforcement, Taxation, Land Planning, Project Monitoring), achieving 50-60% revenue growth',
            'Implemented OCR-based enforcement systems, GIS mapping with Supermap, and IoT sensor integration for real-time monitoring and crime reporting',
            'Transformed legacy code to SOLID principles and MVVM architecture, rolling out 7 major solutions with improved maintainability and scalability',
            'Established on-premise IIS hosting environment, deployed all applications under single domain with ASP.NET Core, Vue.js, Angular, Supermap GIS, and SQL Server',
            'Led team in requirements gathering, architecture design, and stakeholder engagement, representing organization in smart city development conferences'
        ]
    },
    {
        'title': 'Technical Lead / Senior Software Engineer',
        'company': 'Squad Cell',
        'location': 'Lahore, Pakistan',
        'period': 'November 2012 – March 2018',
        'duration': '5 years 5 months',
        'achievements': [
            'Architected and deployed cloud-hosted e-commerce platform from scratch with .NET MVC, Angular, and Swagger UI, managing end-to-end development lifecycle',
            'Developed data analytics framework for automated product recommendations based on purchase frequency, user ratings, and behavior patterns, increasing customer engagement',
            'Designed independent notification microservice (push, in-app, email) serving multiple applications, with reward/discount strategies increasing user retention',
            'Standardized Agile/Scrum practices and automated DevOps branching/build/deployment processes, improving team efficiency and code quality through structured workflows',
            'Mentored team members in OOP principles, design patterns, and algorithms through workshops and code reviews, contributing to company\'s revenue growth and customer base expansion'
        ]
    }
]

# Education
EDUCATION = {
    'degree': 'Bachelor of Computer System Engineering (Honors)',
    'university': 'University of Engineering and Technology',
    'location': 'Lahore, Pakistan',
    'period': '2006 - 2013'
}

# Awards
AWARDS = [
    {
        'title': 'National Speed Programming Champion',
        'organization': 'ACM All Pakistan Speed Programming Competition',
        'year': '2010',
        'description': 'Secured first position competing against 500+ top-tier developers from leading universities nationwide'
    },
    {
        'title': 'HR IT Expert Award',
        'organization': 'Guinness World Records Pakistan - Punjab Youth Festival',
        'year': '2013',
        'description': 'Recognized as premier IT expert for orchestrating technical infrastructure of Pakistan\'s largest youth festival, managing 150,000+ participants'
    }
]

def add_heading_with_style(doc, text, level=1, color=None):
    """Add a heading with custom styling"""
    heading = doc.add_heading(text, level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = color
    return heading

def add_formatted_paragraph(doc, text, bold=False, italic=False, size=None, color=None):
    """Add a formatted paragraph"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return para

def create_header(doc):
    """Create professional header section"""
    # Name
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(PERSONAL_INFO['name'].upper())
    name_run.font.size = Pt(22)
    name_run.bold = True
    name_run.font.color.rgb = RGBColor(5, 99, 187)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(PERSONAL_INFO['title'])
    title_run.font.size = Pt(12)
    title_run.italic = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact Information
    contact_para = doc.add_paragraph()
    contact_text = f"{PERSONAL_INFO['location']} | {PERSONAL_INFO['phone']} | {PERSONAL_INFO['email']}"
    contact_run = contact_para.add_run(contact_text)
    contact_run.font.size = Pt(10)
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Links
    links_para = doc.add_paragraph()
    links_text = f"LinkedIn: {PERSONAL_INFO['linkedin']} | GitHub: {PERSONAL_INFO['github']} | Portfolio: {PERSONAL_INFO['portfolio']}"
    links_run = links_para.add_run(links_text)
    links_run.font.size = Pt(9)
    links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Horizontal line
    doc.add_paragraph('_' * 80)

def create_summary_section(doc, role_type):
    """Create professional summary tailored to role"""
    doc.add_heading('PROFESSIONAL SUMMARY', 1)
    
    if role_type == 'engineer':
        summary = (
            f"Senior Software Engineer with {17}+ years of experience architecting enterprise-scale solutions for "
            "Fortune 500 companies. Azure Cloud & DevOps expert specializing in .NET Core, microservices, and CI/CD automation. "
            "Proven track record: Delivered mission-critical systems processing 2.5M+ data points daily, achieved 65% performance "
            "improvements, and led teams managing billions in operational costs. Expert in AI-assisted development using Cursor AI, "
            "Claude, and ChatGPT to accelerate code generation and improve architectural decision-making. Mentored 50+ developers "
            "across global teams. Strong expertise in full-stack development, cloud architecture, and modern software engineering practices."
        )
    elif role_type == 'architect':
        summary = (
            f"Software Architect with {17}+ years of experience designing and implementing enterprise-scale solutions for "
            "Fortune 500 companies. Azure Cloud Architecture expert specializing in microservices, distributed systems, and scalable "
            "infrastructure. Proven track record: Architected systems processing 2.5M+ data points daily, achieved 65% performance "
            "improvements, and designed cloud solutions reducing infrastructure costs by 40%. Expert in leveraging AI-powered tools "
            "(Cursor AI, Claude, ChatGPT) for architectural design, code generation, and technical decision-making. Led cross-functional "
            "teams of 8+ engineers, mentored 50+ developers, and established technical standards across multiple organizations. "
            "Strong expertise in system design, cloud architecture, design patterns, and modern software engineering principles."
        )
    else:  # managerial
        summary = (
            f"Technical Manager & Lead with {17}+ years of experience leading engineering teams and delivering enterprise-scale solutions "
            "for Fortune 500 companies. Proven leadership track record: Managed distributed teams across 5+ time zones, delivered projects "
            "25% faster than industry standards, and achieved 99.9% delivery success rate. Azure Cloud & DevOps expert with deep technical "
            "expertise in .NET Core, microservices, and CI/CD automation. Expert in leveraging AI development tools (Cursor AI, Claude, ChatGPT) "
            "to enhance team productivity and code quality. Mentored 50+ developers, established Agile/Scrum practices, and standardized "
            "development processes across multiple organizations. Strong expertise in team leadership, technical strategy, stakeholder management, "
            "and driving technical excellence."
        )
    
    para = doc.add_paragraph(summary)
    para_format = para.paragraph_format
    para_format.space_after = Pt(12)

def create_skills_section(doc, role_type):
    """Create comprehensive skills section"""
    doc.add_heading('TECHNICAL SKILLS', 1)
    
    skill_categories = [
        ('Programming Languages', SKILLS['languages']),
        ('Frameworks & Libraries', SKILLS['frameworks']),
        ('Databases', SKILLS['databases']),
        ('Cloud & Azure Services', SKILLS['cloud']),
        ('DevOps & CI/CD', SKILLS['devops']),
        ('Development Tools', SKILLS['tools']),
        ('AI-Assisted Development Tools', SKILLS['ai_tools']),
        ('Methodologies & Patterns', SKILLS['methodologies']),
        ('GIS & Spatial Technologies', SKILLS['gis']),
        ('Monitoring & Analytics', SKILLS['monitoring']),
        ('Additional Skills', SKILLS['other'])
    ]
    
    for category, skills in skill_categories:
        para = doc.add_paragraph()
        run = para.add_run(f"{category}: ")
        run.bold = True
        run.font.size = Pt(11)
        para.add_run(', '.join(skills))
        para_format = para.paragraph_format
        para_format.space_after = Pt(6)

def create_experience_section(doc, role_type):
    """Create work experience section"""
    doc.add_heading('PROFESSIONAL EXPERIENCE', 1)
    
    for exp in EXPERIENCE:
        # Job Title and Company
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(exp['title'])
        title_run.bold = True
        title_run.font.size = Pt(12)
        title_run.font.color.rgb = RGBColor(5, 99, 187)
        company_run = title_para.add_run(f" | {exp['company']}")
        company_run.bold = True
        company_run.font.size = Pt(12)
        
        # Location and Period
        period_para = doc.add_paragraph()
        period_run = period_para.add_run(f"{exp['location']} | {exp['period']} ({exp['duration']})")
        period_run.italic = True
        period_run.font.size = Pt(10)
        
        # Achievements
        for achievement in exp['achievements']:
            para = doc.add_paragraph(achievement, style='List Bullet')
            para_format = para.paragraph_format
            para_format.left_indent = Inches(0.25)
            para_format.space_after = Pt(6)
        
        doc.add_paragraph()  # Space between jobs

def create_education_section(doc):
    """Create education section"""
    doc.add_heading('EDUCATION', 1)
    
    para = doc.add_paragraph()
    degree_run = para.add_run(EDUCATION['degree'])
    degree_run.bold = True
    degree_run.font.size = Pt(11)
    para.add_run(f" | {EDUCATION['university']}, {EDUCATION['location']} | {EDUCATION['period']}")
    
    para_format = para.paragraph_format
    para_format.space_after = Pt(12)

def create_awards_section(doc):
    """Create awards section"""
    doc.add_heading('AWARDS & ACHIEVEMENTS', 1)
    
    for award in AWARDS:
        para = doc.add_paragraph()
        title_run = para.add_run(f"{award['title']} ({award['year']})")
        title_run.bold = True
        title_run.font.size = Pt(11)
        para.add_run(f" | {award['organization']}")
        
        desc_para = doc.add_paragraph(award['description'], style='List Bullet')
        desc_para_format = desc_para.paragraph_format
        desc_para_format.left_indent = Inches(0.25)
        desc_para_format.space_after = Pt(6)

def create_resume(role_type, filename):
    """Create a complete resume document"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Create sections
    create_header(doc)
    doc.add_paragraph()  # Space
    
    create_summary_section(doc, role_type)
    create_skills_section(doc, role_type)
    create_experience_section(doc, role_type)
    create_education_section(doc)
    create_awards_section(doc)
    
    # Save document
    doc.save(filename)
    print(f"✓ Created: {filename}")

def main():
    """Generate all three resume documents"""
    print("Generating ATS-compatible resumes for New Zealand job market...")
    print("=" * 60)
    
    # Create output directory
    output_dir = "resumes"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Generate resumes
    create_resume('engineer', f'{output_dir}/Waqas_Ahmad_Senior_Software_Engineer_Resume.docx')
    create_resume('architect', f'{output_dir}/Waqas_Ahmad_Software_Architect_Resume.docx')
    create_resume('managerial', f'{output_dir}/Waqas_Ahmad_Technical_Manager_Resume.docx')
    
    print("=" * 60)
    print("All resumes generated successfully!")
    print(f"\nFile locations:")
    print(f"1. {os.path.abspath(output_dir)}/Waqas_Ahmad_Senior_Software_Engineer_Resume.docx")
    print(f"2. {os.path.abspath(output_dir)}/Waqas_Ahmad_Software_Architect_Resume.docx")
    print(f"3. {os.path.abspath(output_dir)}/Waqas_Ahmad_Technical_Manager_Resume.docx")

if __name__ == '__main__':
    main()


